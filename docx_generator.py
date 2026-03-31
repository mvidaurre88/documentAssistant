from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import *
from docx.oxml.ns import qn
from word_formatter import format_richtext, agregar_bookmark, agregar_link_interno
import json
import os
from jinja2 import Undefined
import streamlit as st
import subprocess
import tempfile
from io import BytesIO

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# -- Función principal para generar el documento -------------------------------------------------------
def generate_docx(jsonText, document=None, mode=None):    
    
    # -- Cargo el template -----------------------------------------------------------------------------
    if document is not None:
        tpl = DocxTemplate(os.path.join(BASE_DIR, "templates", document + ".docx"))
    else:
        st.error("Error 500: Ocurrió un error con el template del documento")
        return

    # -- Cargo el JSON obtenido ------------------------------------------------------------------------
    if (jsonText is not None):
        context = json.loads(jsonText)
    else:
        st.error("Error 500: Ocurrió un error al obtener la respuesta de la IA")
        return

    # -- Reviso si es PDD o SDD y genero el archivo ----------------------------------------------------
    if document == "SDD":
        buffer = generate_SDD(tpl, context)
    elif document == "PDD":
        buffer = generate_PDD(tpl, context)
    else:
        st.error("Error 500: Tipo de documento no soportado")
        return   
    st.session_state.doc_buffer = buffer 

# -- Funcion para generar SDD --------------------------------------------------------------------------
def generate_SDD(tpl, context):
    
    # Formateo todos los campos que tengan richText
    for tarea in context.get("solucionTecnicaDetallada", []):
        texto_original = tarea.get("descripcionExacta", "")
        # Reemplazamos el string por el objeto RichText
        tarea["descripcionExacta"] = format_richtext(texto_original)
        
    # Formateo ejecucion y reejecucion
    for campo in ["ejecucion", "reejecucion"]:
        if campo in context and isinstance(context[campo], str):
            context[campo] = format_richtext(context[campo])
    
    # Genero la imagen del diagrama de pasos
    generar_imagen(field="diagrama_pasos", height=7.87, tpl=tpl, context=context)

    # Genero la imagen del diagrama de bajo nivel
    generar_imagen(field="diagrama_detalle", height=8.7, tpl=tpl, context=context)
    
    tpl.render(context)

    buffer = BytesIO()
    tpl.save(buffer)
    buffer.seek(0)

    return buffer

def generate_PDD(tpl, context):
    # Formateo todos los campos que tengan richText
    campos_richtext = ["propositoProceso"]
    for campo in campos_richtext:
        if campo in context and isinstance(context[campo], str):
            context[campo] = format_richtext(context[campo])

    # Formateo las excepciones
    excepciones = context.get("excepciones", [])
    
    excepcionesSys = [e for e in excepciones if e["tipo"] in ("tecnica", "técnica")]
    context["excepcionesSys"] = excepcionesSys
    
    excepcionesNeg = [e for e in excepciones if e["tipo"] == "negocio"]
    context["excepcionesNeg"] = excepcionesNeg

    mapa_excepciones = {e["escenario"]: e["numero"] for e in excepciones}
    fases_data = context["fases"]
    
    # Enumerar pasos y agregar número de excepción desde el mapa
    paso_counter = [1]
    for fase in fases_data:
        for paso in fase["pasos"]:
            paso["numero"] = f"3.1.{paso_counter[0]}"
            paso["excepcion_numero"] = mapa_excepciones.get(paso["excepcion_escenario"], "")
            paso_counter[0] += 1

    # Generar subdocs DESPUÉS de crear tpl
    fases_subdocs = [crear_subdoc_fase(tpl, f) for f in fases_data]

    variables_template = tpl.get_undeclared_template_variables()

    context["fases"] = fases_data

    faltantes = [var for var in variables_template if var not in context]
    vacios     = [key for key, value in context.items() if value == "" or value is None]
    print("Campos faltantes:", faltantes)
    print("Campos vacios:", vacios)

    variables_template = tpl.get_undeclared_template_variables()
    for var in variables_template:
        val = context.get(var)
        if val is None or val == "":
            context[var] = f"{{{{{var}}}}}"

    tpl.render(context)
    
    buffer = BytesIO()
    tpl.save(buffer)
    buffer.seek(0)

    doc = Document(buffer)

    # Insertar tabla después de cada título de fase
    for fase in fases_data:
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == fase["tituloFase"]:
                tabla = doc.add_table(rows=1, cols=4)
                tabla.style = 'Table Grid'
                headers = ["N°", "Acción", "Detalle", "Excepción"]
                for j, h in enumerate(headers):
                    cell = tabla.rows[0].cells[j]
                    cell.text = h
                    # Fondo naranja
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'ED7D31')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    tcPr.append(shd)
                    # Texto blanco y negrita
                    run = cell.paragraphs[0].runs[0]
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                for paso in fase["pasos"]:
                    row = tabla.add_row()
                    row.cells[0].text = paso["numero"]
                    row.cells[1].text = paso["accion"]
                    row.cells[2].text = paso["detalle"]
                    row.cells[3].text = paso.get("excepcion_numero", "")
                para._p.addnext(tabla._tbl)
                break
    
    WIDTHS_FASES        = [Inches(0.6), Inches(2.0), Inches(3.3), Inches(0.9)]
    WIDTHS_EXCEPCIONES  = [Inches(0.6), Inches(2.4), Inches(3.7)]
    
    for tabla in doc.tables:
        primera_fila = [c.text.strip() for c in tabla.rows[0].cells] if tabla.rows else []

        if "Acción" in primera_fila and "Detalle" in primera_fila:
            set_col_widths(tabla, WIDTHS_FASES)

        elif "Escenario" in primera_fila or "Excepción" in primera_fila:
            set_col_widths(tabla, WIDTHS_EXCEPCIONES)
    
    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    return output_buffer
    
class KeepUndefined(Undefined):
    def __str__(self):
        return f"{{{{{self._undefined_name}}}}}"
    def __repr__(self):
        return f"{{{{{self._undefined_name}}}}}"

def crear_subdoc_fase(tpl, fase):
    """Crea un subdocumento con título + tabla para una fase."""
    sd = tpl.new_subdoc()
    sd.add_paragraph(fase["tituloFase"], style="Heading 2")
    tabla = sd.add_table(rows=1, cols=4)
    set_col_widths(tabla, [Cm(0.6), Cm(2.0), Cm(3.3), Cm(0.9)])
    tabla.style = 'Table Grid'
    headers = ["N°", "Acción", "Detalle", "Excepción"]
    for i, h in enumerate(headers):
        tabla.rows[0].cells[i].text = h
    for paso in fase["pasos"]:
        row = tabla.add_row()
        #row.cells[0].text = paso["numero"]
        row.cells[1].text = paso["accion"]
        row.cells[2].text = paso["detalle"]
        #row.cells[3].text = paso.get("excepcion_numero", "")
    return sd

def set_col_widths(tabla, widths):
    for row in tabla.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            
def generar_imagen(field=None, height=5, tpl=None, context=None):
    mermaidCode = context.get(field)

    with tempfile.TemporaryDirectory() as tmpdir:
        mmd_path = f"{tmpdir}/diagrama.mmd"
        png_path = f"{tmpdir}/diagrama.png"

        # Crear .mmd temporal
        with open(mmd_path, "w", encoding="utf-8") as f:
            f.write("%%{init: {\"theme\": \"neutral\", \"flowchart\": {\"curve\": \"stepAfter\", \"rankSpacing\": 40}}}%%\n" + mermaidCode + "\nlinkStyle default stroke-width:4px;")

        # Generar imagen
        subprocess.run(
            f"mmdc -i {mmd_path} -o {png_path} -w 1200 -H 600",
            shell=True,
            check=True
        )

        with open(png_path, "rb") as f:
            image_stream = BytesIO(f.read())
            
    # Insertar imagen
    image = InlineImage(tpl, image_descriptor=image_stream, height=Inches(height))
    context[field] = image
    
