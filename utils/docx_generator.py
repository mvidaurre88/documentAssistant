from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import *
from docx.oxml.ns import qn
from utils.word_formatter import format_richtext, agregar_bookmark, agregar_link_interno
import json
import os
from jinja2 import Undefined
import streamlit as st
import subprocess
import tempfile
from io import BytesIO

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# -- Función principal para generar el documento -------------------------------------------------------
def generate_docx(jsonText: dict, document=None, mode=None):    
    
    # -- Cargo el template -----------------------------------------------------------------------------
    if document is not None:
        tpl = DocxTemplate(os.path.join(BASE_DIR, "templates", document + ".docx"))
    else:
        st.error("Error 500: Ocurrió un error con el template del documento")
        return

    # -- Reviso si es PDD o SDD y genero el archivo ----------------------------------------------------
    jsonText = sanitize_all(jsonText)
    if document == "SDD":
        buffer = generate_SDD(tpl, jsonText)
    elif document == "PDD":
        buffer = generate_PDD(tpl, jsonText)
    else:
        st.error("Error 500: Tipo de documento no soportado")
        return   
    st.session_state.doc_buffer = buffer 

# -- Funcion para generar SDD --------------------------------------------------------------------------
def generate_SDD(tpl, context):
    
    # Formateo todos los campos que tengan richText
    for tarea in context.get("solucionTecnicaDetallada", []):
        texto_original = tarea.get("descripcionExacta", "")
        # Reemplazamos el string por el objeto RichText
        tarea["descripcionExacta"] = format_richtext(texto_original)
        
    # Formateo ejecucion y reejecucion
    for campo in ["ejecucion", "reejecucion"]:
        if campo in context and isinstance(context[campo], str):
            context[campo] = format_richtext(context[campo])
    
    # GENERO IMAGENES
    load_img(field="diagrama_pasos", height=7.87, tpl=tpl, context=context)
    load_img(field="diagrama_detalle", height=8.7, tpl=tpl, context=context)
    
    tpl.render(context)

    buffer = BytesIO()
    tpl.save(buffer)
    buffer.seek(0)

    return buffer

def generate_PDD(tpl, context: dict):
    
    # PROPOSITO PROCESO ----------------------------------------------------
    context["propositoProceso"] = format_richtext(context["propositoProceso"])

    # IMAGENES DE LOS DIAGRAMAS --------------------------------------------
    load_img(field="diagramaAltoNivel", height=6, tpl=tpl, context=context)
    load_img(field="diagramaBajoNivel", height=8, tpl=tpl, context=context)

    # PASOS DE LAS FASES ------------------------------------------------------
    fases = context.get("fases", [])
    phaseCounter = 1
    for fase in fases:
        stepCounter = 1
        pasos = fase.get("pasos", [])
        for paso in pasos:
            paso["numero"] = f"3.{phaseCounter}.{stepCounter}"
            stepCounter += 1
        phaseCounter += 1

    # EXCEPCIONES -------------------------------------------------------------
    excepciones = context.get("excepciones", [])
    excepcionesSys = []
    excepcionesNeg = []
    counter_exceptions = 1
    for e in excepciones:
        if e["tipo"].strip().lower() in ("tecnica", "técnica"):
            e["numero"] = f"4.2.{counter_exceptions}"
            excepcionesSys.append(e)
            counter_exceptions += 1
    for e in excepciones:
        if e["tipo"].strip().lower() in ("negocio", "de negocio"):
            e["numero"] = f"4.2.{counter_exceptions}"
            excepcionesNeg.append(e)
            counter_exceptions += 1
    context["excepcionesSys"] = excepcionesSys
    context["excepcionesNeg"] = excepcionesNeg
    
    variables_template = tpl.get_undeclared_template_variables()

    faltantes = [var for var in variables_template if var not in context]
    vacios     = [key for key, value in context.items() if value == "" or value is None]
    print("Campos faltantes:", faltantes)
    print("Campos vacios:", vacios)

    variables_template = tpl.get_undeclared_template_variables()
    for var in variables_template:
        val = context.get(var)
        if val is None or val == "":
            context[var] = f"{{{{{var}}}}}"

    tpl.render(context)
    
    buffer = BytesIO()
    tpl.save(buffer)
    buffer.seek(0)

    doc = Document(buffer)
    
    WIDTHS_PRECONDICIONES  = [Inches(0.6), Inches(1.4), Inches(4.7)]
    WIDTHS_FASES        = [Inches(0.6), Inches(1.5), Inches(3.2), Inches(1.4)]
    WIDTHS_EXCEPCIONES  = [Inches(0.6), Inches(2.4), Inches(3.7)]
    
    for tabla in doc.tables:
        primera_fila = [c.text.strip() for c in tabla.rows[0].cells] if tabla.rows else []

        if "N°" in primera_fila and "Precondición" in primera_fila and "Detalle" in primera_fila:
            set_col_widths(tabla, WIDTHS_PRECONDICIONES)

        elif "Acción" in primera_fila and "Detalle" in primera_fila:
            set_col_widths(tabla, WIDTHS_FASES)

        elif "N°" in primera_fila and "Escenario" in primera_fila and "Acción" in primera_fila:
            set_col_widths(tabla, WIDTHS_EXCEPCIONES)
    
    output_buffer = BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    
    return output_buffer

def crear_subdoc_fase(tpl, fase):
    """Crea un subdocumento con título + tabla para una fase."""
    sd = tpl.new_subdoc()
    sd.add_paragraph(fase["tituloFase"], style="Heading 2")
    tabla = sd.add_table(rows=1, cols=4)
    set_col_widths(tabla, [Cm(0.6), Cm(2.0), Cm(3.3), Cm(0.9)])
    tabla.style = 'Table Grid'
    headers = ["N°", "Acción", "Detalle", "Excepción"]
    for i, h in enumerate(headers):
        tabla.rows[0].cells[i].text = h
    for paso in fase["pasos"]:
        row = tabla.add_row()
        #row.cells[0].text = paso["numero"]
        row.cells[1].text = paso["accion"]
        row.cells[2].text = paso["detalle"]
        #row.cells[3].text = paso.get("excepcion_numero", "")
    return sd
            
def load_img(field=None, height=5, tpl=None, context=None):
    image = st.session_state.get(field, None)
    diagram = ""
    
    if image:
        try:
            image_stream = BytesIO(image)
            image_stream.seek(0)
            diagram = InlineImage(tpl, image_descriptor=image_stream, height=Inches(height))
        except Exception as e:
            diagram = "[Error al cargar imagen: " + str(e) + "]"
        
    context[field] = diagram

# ELIMINA CARACTERES < y > para evitar problemas con el XML del docx ---------------------------------
def sanitize_all(data):
    if isinstance(data, dict):
        return {k: sanitize_all(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [sanitize_all(i) for i in data]
    elif isinstance(data, str):
        return data.replace("<", "&lt;").replace(">", "&gt;")
    return data

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

def set_col_widths(table, widths):
    """
    Fuerza el ancho de columnas en una tabla word-docx.
    Desactiva autofit y opera directamente sobre XML para evitar
    que Word expanda celdas con texto largo sin espacios.
    """
    # 1. Forzar layout fijo (desactiva autofit)
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

    # 2. Setear ancho total de la tabla
    total_emu = sum(w.inches * 914400 for w in widths)
    total_twips = int(sum(w.inches * 1440 for w in widths))

    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(total_twips))
    tblW.set(qn('w:type'), 'dxa')

    # 3. Forzar ancho en cada celda de cada fila
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(widths):
                break
            width_twips = int(widths[i].inches * 1440)
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.insert(0, tcW)
            tcW.set(qn('w:w'), str(width_twips))
            tcW.set(qn('w:type'), 'dxa')

            # 4. Activar word wrap en cada párrafo de la celda
            for para in cell.paragraphs:
                pPr = para._p.find(qn('w:pPr'))
                if pPr is None:
                    pPr = OxmlElement('w:pPr')
                    para._p.insert(0, pPr)
                # Desactivar "no wrap"
                wordWrap = pPr.find(qn('w:wordWrap'))
                if wordWrap is not None:
                    pPr.remove(wordWrap)